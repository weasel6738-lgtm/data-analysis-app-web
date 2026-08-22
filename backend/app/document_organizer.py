"""Offline document extraction, classification, and safe archive creation."""

from __future__ import annotations

import csv
import io
import re
import zipfile
from collections import Counter, defaultdict
from dataclasses import dataclass
from pathlib import PurePosixPath
from typing import Iterable, Mapping, Sequence

from docx import Document
from openpyxl import load_workbook
from pypdf import PdfReader


CATEGORY_KEYWORDS: dict[str, tuple[str, ...]] = {
    "계약서": (
        "계약서",
        "계약",
        "협약",
        "합의서",
        "agreement",
        "contract",
        "nda",
        "계약 당사자",
        "비밀유지",
    ),
    "영수증·세금계산서": (
        "영수증",
        "세금계산서",
        "청구서",
        "공급가액",
        "부가세",
        "결제",
        "receipt",
        "invoice",
    ),
    "보고서": (
        "보고서",
        "분석 결과",
        "조사 결과",
        "월간 보고",
        "주간 보고",
        "결과 보고",
        "report",
    ),
    "회의록": (
        "회의록",
        "회의 일시",
        "참석자",
        "회의 안건",
        "의결",
        "minutes",
        "meeting",
    ),
    "데이터": (
        "데이터",
        "dataset",
        "data",
        "통계",
        "측정값",
        "레코드",
    ),
    "이미지": (),
    "기타/분류불가": (),
}

CATEGORIES: tuple[str, ...] = tuple(CATEGORY_KEYWORDS)
TEXT_EXTENSIONS = {".txt", ".md"}
IMAGE_EXTENSIONS = {".bmp", ".gif", ".jpeg", ".jpg", ".png", ".tif", ".tiff", ".webp"}
SUPPORTED_EXTENSIONS = TEXT_EXTENSIONS | {".csv", ".xlsx", ".docx", ".pdf"} | IMAGE_EXTENSIONS
MAX_EXTRACTED_CHARACTERS = 500_000


@dataclass(frozen=True)
class InputDocument:
    """A browser-uploaded file represented entirely in memory."""

    path: str
    content: bytes


@dataclass(frozen=True)
class ExtractionResult:
    text: str
    status: str
    error: str = ""


@dataclass(frozen=True)
class ClassificationResult:
    category: str
    confidence: float
    reason: str
    matched_keywords: tuple[str, ...] = ()


@dataclass
class ProcessedDocument:
    document_id: str
    original_path: str
    original_name: str
    content: bytes
    suggested_category: str
    confidence: float
    reason: str
    matched_keywords: tuple[str, ...]
    processing_status: str
    error: str = ""
    note: str = ""


def _extension(filename: str) -> str:
    name = filename.rsplit("/", 1)[-1].rsplit("\\", 1)[-1]
    dot = name.rfind(".")
    return name[dot:].lower() if dot >= 0 else ""


def normalize_relative_path(path: str) -> str:
    """Return a normalized upload path or reject an unsafe path."""

    if not path or "\x00" in path:
        raise ValueError("파일 경로가 비어 있거나 올바르지 않습니다.")
    normalized = path.replace("\\", "/")
    if normalized.startswith("/") or re.match(r"^[A-Za-z]:", normalized):
        raise ValueError("절대 경로는 업로드 경로로 사용할 수 없습니다.")
    parts = normalized.split("/")
    if any(part in {"", ".", ".."} for part in parts):
        raise ValueError("상위 경로 이동 또는 빈 경로 구간은 허용되지 않습니다.")
    return "/".join(parts)


def sanitize_filename(filename: str) -> str:
    """Create a portable, non-traversing archive filename."""

    name = filename.replace("\\", "/").rsplit("/", 1)[-1]
    name = re.sub(r"[\x00-\x1f<>:\"/\\|?*]", "_", name).strip(" .")
    if not name:
        name = "unnamed"
    stem = name.split(".", 1)[0].upper()
    if stem in {"CON", "PRN", "AUX", "NUL"} or re.fullmatch(r"(COM|LPT)[1-9]", stem):
        name = f"_{name}"
    if len(name) > 180:
        dot = name.rfind(".")
        suffix = name[dot:] if 0 < dot and len(name) - dot <= 20 else ""
        name = f"{name[: 180 - len(suffix)]}{suffix}"
    return name


def _decode_text(content: bytes) -> str:
    errors: list[str] = []
    for encoding in ("utf-8-sig", "cp949"):
        try:
            return content.decode(encoding)
        except UnicodeDecodeError as exc:
            errors.append(f"{encoding}: {exc}")
    raise ValueError("UTF-8 또는 CP949 텍스트로 읽을 수 없습니다. " + " / ".join(errors))


def _trim(text: str) -> str:
    return text[:MAX_EXTRACTED_CHARACTERS]


def extract_text(filename: str, content: bytes) -> ExtractionResult:
    """Extract readable text without allowing one bad file to abort a batch."""

    extension = _extension(filename)
    if extension not in SUPPORTED_EXTENSIONS:
        return ExtractionResult(
            "",
            "unsupported",
            f"지원하지 않는 파일 형식입니다: {extension or '확장자 없음'}",
        )
    if not content:
        return ExtractionResult("", "empty", "파일이 비어 있습니다.")
    if extension in IMAGE_EXTENSIONS:
        return ExtractionResult("", "ready")

    try:
        if extension in TEXT_EXTENSIONS:
            text = _decode_text(content)
        elif extension == ".csv":
            decoded = _decode_text(content)
            rows = csv.reader(io.StringIO(decoded))
            text = "\n".join(" | ".join(row) for row in rows)
        elif extension == ".xlsx":
            workbook = load_workbook(io.BytesIO(content), read_only=True, data_only=True)
            try:
                values: list[str] = []
                extracted_characters = 0
                for sheet in workbook.worksheets:
                    values.append(sheet.title)
                    extracted_characters += len(sheet.title)
                    for row in sheet.iter_rows(values_only=True):
                        line = " | ".join("" if cell is None else str(cell) for cell in row)
                        values.append(line)
                        extracted_characters += len(line)
                        if extracted_characters >= MAX_EXTRACTED_CHARACTERS:
                            break
                    if extracted_characters >= MAX_EXTRACTED_CHARACTERS:
                        break
                text = "\n".join(values)
            finally:
                workbook.close()
        elif extension == ".docx":
            document = Document(io.BytesIO(content))
            paragraphs = [paragraph.text for paragraph in document.paragraphs]
            table_rows = [
                " | ".join(cell.text for cell in row.cells)
                for table in document.tables
                for row in table.rows
            ]
            text = "\n".join(paragraphs + table_rows)
        else:
            reader = PdfReader(io.BytesIO(content), strict=False)
            if reader.is_encrypted:
                return ExtractionResult("", "error", "암호화된 PDF는 내용을 읽을 수 없습니다.")
            text = "\n".join(page.extract_text() or "" for page in reader.pages)
    except Exception as exc:
        return ExtractionResult("", "error", f"파일 내용 읽기 실패: {exc}")

    text = _trim(text).strip()
    if not text:
        return ExtractionResult("", "empty", "추출 가능한 텍스트가 없습니다.")
    return ExtractionResult(text, "ready")


def _normalized_keywords(
    keywords: Mapping[str, Sequence[str]] | None,
) -> dict[str, tuple[str, ...]]:
    source = keywords or CATEGORY_KEYWORDS
    result: dict[str, tuple[str, ...]] = {}
    for category in CATEGORIES:
        values = source.get(category, CATEGORY_KEYWORDS[category])
        result[category] = tuple(
            dict.fromkeys(value.strip().lower() for value in values if value.strip())
        )
    return result


def classify_document(
    filename: str,
    text: str,
    keywords: Mapping[str, Sequence[str]] | None = None,
) -> ClassificationResult:
    """Classify deterministically from filename, extension, and extracted text."""

    extension = _extension(filename)
    if extension in IMAGE_EXTENSIONS:
        return ClassificationResult("이미지", 0.99, f"이미지 확장자({extension})")

    lowered_name = filename.lower()
    lowered_text = text.lower()
    configured = _normalized_keywords(keywords)
    scores: Counter[str] = Counter()
    evidence: defaultdict[str, list[str]] = defaultdict(list)

    for category, category_keywords in configured.items():
        for keyword in category_keywords:
            in_name = keyword in lowered_name
            in_text = keyword in lowered_text
            if in_name:
                scores[category] += 3
                evidence[category].append(f"파일명:{keyword}")
            if in_text:
                scores[category] += 1
                evidence[category].append(f"내용:{keyword}")

    if extension in {".csv", ".xlsx"}:
        scores["데이터"] += 2
        evidence["데이터"].append(f"형식:{extension}")

    ranked = sorted(
        ((scores[category], -index, category) for index, category in enumerate(CATEGORIES)),
        reverse=True,
    )
    winning_score, _, category = ranked[0]
    if winning_score <= 0 or category in {"이미지", "기타/분류불가"}:
        return ClassificationResult(
            "기타/분류불가",
            0.0,
            "일치하는 분류 키워드가 없습니다.",
        )

    runner_up = ranked[1][0]
    confidence = min(0.99, 0.45 + (winning_score / (winning_score + runner_up + 2)) * 0.5)
    matched = tuple(dict.fromkeys(item.split(":", 1)[1] for item in evidence[category]))
    reason = ", ".join(evidence[category][:6])
    return ClassificationResult(category, round(confidence, 2), reason, matched)


def process_documents(
    documents: Iterable[InputDocument],
    keywords: Mapping[str, Sequence[str]] | None = None,
) -> list[ProcessedDocument]:
    """Process every document independently and retain explicit failure states."""

    processed: list[ProcessedDocument] = []
    names: defaultdict[str, list[ProcessedDocument]] = defaultdict(list)
    for index, document in enumerate(documents):
        raw_name = document.path.replace("\\", "/").rsplit("/", 1)[-1]
        safe_name = sanitize_filename(raw_name)
        try:
            original_path = normalize_relative_path(document.path)
            extraction = extract_text(original_path, document.content)
            classification = classify_document(original_path, extraction.text, keywords)
        except ValueError as exc:
            original_path = document.path
            extraction = ExtractionResult("", "error", f"안전하지 않은 업로드 경로: {exc}")
            classification = ClassificationResult(
                "기타/분류불가", 0.0, "경로 안전성 검사 실패"
            )

        item = ProcessedDocument(
            document_id=str(index),
            original_path=original_path,
            original_name=safe_name,
            content=document.content,
            suggested_category=classification.category,
            confidence=classification.confidence,
            reason=classification.reason,
            matched_keywords=classification.matched_keywords,
            processing_status=extraction.status,
            error=extraction.error,
        )
        processed.append(item)
        names[safe_name.casefold()].append(item)

    for duplicates in names.values():
        if len(duplicates) > 1:
            note = f"같은 파일명 {len(duplicates)}개: ZIP 생성 시 이름을 안전하게 변경합니다."
            for item in duplicates:
                item.note = note
    return processed


def _assert_safe_archive_path(path: str) -> None:
    pure_path = PurePosixPath(path)
    if pure_path.is_absolute() or any(part in {"", ".", ".."} for part in pure_path.parts):
        raise ValueError(f"안전하지 않은 ZIP 내부 경로입니다: {path}")


def _deduplicate_name(category: str, filename: str, used: set[str]) -> str:
    stem, dot, suffix = filename.rpartition(".")
    if not dot:
        stem, suffix = filename, ""
    candidate = filename
    counter = 2
    while f"{category}/{candidate}".casefold() in used:
        candidate = f"{stem} ({counter}){'.' + suffix if suffix else ''}"
        counter += 1
    used.add(f"{category}/{candidate}".casefold())
    return candidate


def _category_directory(category: str) -> str:
    return category.replace("/", "_").replace("\\", "_")


def _csv_safe(value: object) -> object:
    if isinstance(value, str) and value.startswith(("=", "+", "-", "@")):
        return f"'{value}"
    return value


def create_organized_zip(
    documents: Sequence[ProcessedDocument],
    overrides: Mapping[str, str] | None = None,
) -> bytes:
    """Create an in-memory ZIP containing originals and an audit manifest."""

    final_categories = dict(overrides or {})
    unknown_ids = set(final_categories) - {document.document_id for document in documents}
    if unknown_ids:
        raise ValueError(f"알 수 없는 문서 ID가 있습니다: {', '.join(sorted(unknown_ids))}")

    output = io.BytesIO()
    manifest = io.StringIO(newline="")
    fieldnames = [
        "원본 경로",
        "원본 파일명",
        "추천 분류",
        "최종 분류",
        "신뢰도",
        "분류 근거",
        "일치 키워드",
        "처리 상태",
        "오류",
        "참고",
        "ZIP 내부 경로",
    ]
    writer = csv.DictWriter(manifest, fieldnames=fieldnames)
    writer.writeheader()
    used: set[str] = set()

    with zipfile.ZipFile(output, "w", compression=zipfile.ZIP_DEFLATED) as archive:
        for document in documents:
            category = final_categories.get(document.document_id, document.suggested_category)
            if category not in CATEGORIES:
                raise ValueError(f"허용되지 않은 분류입니다: {category}")
            category_directory = _category_directory(category)
            filename = _deduplicate_name(
                category_directory, sanitize_filename(document.original_name), used
            )
            archive_path = f"{category_directory}/{filename}"
            _assert_safe_archive_path(archive_path)
            archive.writestr(archive_path, document.content)
            writer.writerow(
                {
                    "원본 경로": _csv_safe(document.original_path),
                    "원본 파일명": _csv_safe(document.original_name),
                    "추천 분류": document.suggested_category,
                    "최종 분류": category,
                    "신뢰도": f"{document.confidence:.2f}",
                    "분류 근거": document.reason,
                    "일치 키워드": ", ".join(document.matched_keywords),
                    "처리 상태": document.processing_status,
                    "오류": document.error,
                    "참고": document.note,
                    "ZIP 내부 경로": archive_path,
                }
            )
        archive.writestr("manifest.csv", "\ufeff" + manifest.getvalue())
    return output.getvalue()
