import csv
import io
import zipfile

import pytest
from docx import Document
from openpyxl import Workbook
from pypdf import PdfWriter

try:
    from app.document_organizer import (
        InputDocument,
        classify_document,
        create_organized_zip,
        extract_text,
        normalize_relative_path,
        process_documents,
    )
except ModuleNotFoundError:
    from backend.app.document_organizer import (
        InputDocument,
        classify_document,
        create_organized_zip,
        extract_text,
        normalize_relative_path,
        process_documents,
    )


def _xlsx_bytes() -> bytes:
    output = io.BytesIO()
    workbook = Workbook()
    workbook.active.append(["측정값", 42])
    workbook.save(output)
    workbook.close()
    return output.getvalue()


def _docx_bytes() -> bytes:
    output = io.BytesIO()
    document = Document()
    document.add_paragraph("주간 보고서 분석 결과")
    document.save(output)
    return output.getvalue()


def _encrypted_pdf_bytes() -> bytes:
    output = io.BytesIO()
    writer = PdfWriter()
    writer.add_blank_page(width=100, height=100)
    writer.encrypt("secret")
    writer.write(output)
    return output.getvalue()


def test_extracts_representative_supported_documents():
    assert "계약" in extract_text("contract.txt", "계약 조건".encode()).text
    assert "참석자" in extract_text("minutes.csv", "항목,값\n참석자,홍길동".encode()).text
    assert "측정값" in extract_text("measurements.xlsx", _xlsx_bytes()).text
    assert "보고서" in extract_text("weekly.docx", _docx_bytes()).text


def test_classification_is_deterministic_when_content_order_changes():
    first = classify_document("notes.txt", "참석자 홍길동 회의 안건 예산 회의록")
    shuffled = classify_document("notes.txt", "회의록 예산 회의 안건 홍길동 참석자")

    assert first.category == shuffled.category == "회의록"
    assert first.confidence == shuffled.confidence
    assert set(first.matched_keywords) == set(shuffled.matched_keywords)


def test_bad_empty_unsupported_and_encrypted_files_are_isolated():
    records = process_documents(
        [
            InputDocument("empty.txt", b""),
            InputDocument("program.exe", b"MZ"),
            InputDocument("broken.docx", b"not a zip"),
            InputDocument("secret.pdf", _encrypted_pdf_bytes()),
        ]
    )

    assert [record.processing_status for record in records] == [
        "empty",
        "unsupported",
        "error",
        "error",
    ]
    assert "지원하지 않는" in records[1].error
    assert "읽기 실패" in records[2].error
    assert "암호화" in records[3].error


def test_duplicate_names_overrides_and_manifest_are_preserved():
    records = process_documents(
        [
            InputDocument("첫번째/report.txt", "보고서 하나".encode()),
            InputDocument("두번째/report.txt", "보고서 둘".encode()),
        ]
    )
    archive_bytes = create_organized_zip(
        records, {record.document_id: "계약서" for record in records}
    )

    with zipfile.ZipFile(io.BytesIO(archive_bytes)) as archive:
        assert "계약서/report.txt" in archive.namelist()
        assert "계약서/report (2).txt" in archive.namelist()
        assert archive.read("계약서/report.txt") == "보고서 하나".encode()
        manifest = archive.read("manifest.csv").decode("utf-8-sig")

    rows = list(csv.DictReader(io.StringIO(manifest)))
    assert [row["최종 분류"] for row in rows] == ["계약서", "계약서"]
    assert all("같은 파일명" in row["참고"] for row in rows)
    assert rows[1]["원본 경로"] == "두번째/report.txt"


def test_path_traversal_is_rejected_and_never_written_to_zip():
    with pytest.raises(ValueError):
        normalize_relative_path("../outside.txt")

    records = process_documents([InputDocument("../outside.txt", b"secret")])
    assert records[0].processing_status == "error"
    assert "안전하지 않은" in records[0].error

    with zipfile.ZipFile(io.BytesIO(create_organized_zip(records))) as archive:
        names = archive.namelist()
        assert "기타_분류불가/outside.txt" in names
        assert all(not name.startswith("/") and ".." not in name.split("/") for name in names)

    with pytest.raises(ValueError):
        create_organized_zip(records, {"0": "../escape"})
